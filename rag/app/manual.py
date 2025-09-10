#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import copy
import re

from api.db import ParserType, LLMType
from api.db.services.llm_service import LLMBundle
from io import BytesIO
from rag.nlp import rag_tokenizer, tokenize, tokenize_table, bullets_category, title_frequency, tokenize_chunks, docx_question_level
from rag.utils import num_tokens_from_string
from deepdoc.parser import PdfParser, PlainParser, DocxParser
from deepdoc.parser.pdf_parser import VisionParser
from docx import Document
from PIL import Image


class Pdf(PdfParser):
    def __init__(self):
        self.model_speciess = ParserType.MANUAL.value
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, zoomin=3, callback=None):
        from timeit import default_timer as timer

        start = timer()
        callback(msg="OCR started")
        self.__images__(filename if not binary else binary, zoomin, from_page, to_page, callback)
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.debug("OCR: {}".format(timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.65, "Layout analysis ({:.2f}s)".format(timer() - start))
        logging.debug("layouts: {}".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.67, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        tbls = self._extract_table_figure(True, zoomin, True, True)
        self._concat_downward()
        self._filter_forpages()
        callback(0.68, "Text merged ({:.2f}s)".format(timer() - start))

        # clean mess
        for b in self.boxes:
            b["text"] = re.sub(r"([\t 　]|\u3000){2,}", " ", b["text"].strip())

        return [(b["text"], b.get("layoutno", ""), self.get_position(b, zoomin)) for i, b in enumerate(self.boxes)], tbls


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath(".//pic:pic")
        if not img:
            return None
        img = img[0]
        embed = img.xpath(".//a:blip/@r:embed")[0]
        related_part = document.part.related_parts[embed]
        image = related_part.image
        image = Image.open(BytesIO(image.blob))
        return image

    def concat_img(self, img1, img2):
        if img1 and not img2:
            return img1
        if not img1 and img2:
            return img2
        if not img1 and not img2:
            return None
        width1, height1 = img1.size
        width2, height2 = img2.size

        new_width = max(width1, width2)
        new_height = height1 + height2
        new_image = Image.new("RGB", (new_width, new_height))

        new_image.paste(img1, (0, 0))
        new_image.paste(img2, (0, height1))

        return new_image

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, callback=None, vision_model=None, vision_prompt=None):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        pn = 0
        last_answer, last_image = "", None
        question_stack, level_stack = [], []
        ti_list = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = 0, ""
            if from_page <= pn < to_page and p.text.strip():
                question_level, p_text = docx_question_level(p)
            if not question_level or question_level > 6:  # not a question
                # Process images with vision model if available
                current_image = self.get_picture(self.doc, p)
                if current_image and vision_model and vision_prompt:
                    try:
                        image_text = vision_model.describe_with_prompt(current_image, vision_prompt)
                        last_answer = f"{last_answer}\n{p_text}\n[Image Content]: {image_text}"
                    except Exception as e:
                        logging.warning(f"Vision model error: {e}. Using original image.")
                        last_answer = f"{last_answer}\n{p_text}"
                        last_image = self.concat_img(last_image, current_image)
                else:
                    last_answer = f"{last_answer}\n{p_text}"
                    last_image = self.concat_img(last_image, current_image)
            else:  # is a question
                if last_answer or last_image:
                    sum_question = "\n".join(question_stack)
                    if sum_question:
                        ti_list.append((f"{sum_question}\n{last_answer}", last_image))
                    last_answer, last_image = "", None

                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(p_text)
                level_stack.append(question_level)
            for run in p.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    pn += 1
                    continue
                if "w:br" in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        if last_answer:
            sum_question = "\n".join(question_stack)
            if sum_question:
                ti_list.append((f"{sum_question}\n{last_answer}", last_image))

        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                        else:
                            break
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return ti_list, tbls


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
    PDF and DOCX are supported, with optional vision model for images.
    """
    parser_config = kwargs.get("parser_config", {"chunk_token_num": 512, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    pdf_parser = None
    doc = {"docnm_kwd": filename}
    doc["title_tks"] = rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", doc["docnm_kwd"]))
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    # is it English
    eng = lang.lower() == "english"  # pdf_parser.is_english

    if re.search(r"\.pdf$", filename, re.IGNORECASE):
        layout_recognizer = parser_config.get("layout_recognize", "DeepDOC")

        # Handle different layout recognizers
        if layout_recognizer == "Plain Text":
            pdf_parser = PlainParser()
            sections, tbls = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback)
        elif layout_recognizer == "DeepDOC":
            pdf_parser = Pdf()
            sections, tbls = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback)
        else:
            # Use Vision model (like Gemini) for PDF parsing
            try:
                vision_model = LLMBundle(kwargs["tenant_id"], LLMType.IMAGE2TEXT, llm_name=layout_recognizer, lang=lang)
                pdf_parser = VisionParser(vision_model=vision_model, **kwargs)
                sections, tbls = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback)
                callback(0.8, "Vision model parsing completed.")
            except Exception as e:
                logging.warning(f"Failed to use vision model {layout_recognizer}: {e}. Falling back to DeepDOC.")
                pdf_parser = Pdf()
                sections, tbls = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback)
        if sections and len(sections[0]) < 3:
            sections = [(t, lvl, [[0] * 5]) for t, lvl in sections]
        # set pivot using the most frequent type of title,
        # then merge between 2 pivot
        # Check if pdf_parser has outlines attribute (not available in VisionParser)
        if hasattr(pdf_parser, "outlines") and len(sections) > 0 and len(pdf_parser.outlines) / len(sections) > 0.03:
            max_lvl = max([lvl for _, lvl in pdf_parser.outlines])
            most_level = max(0, max_lvl - 1)
            levels = []
            for txt, _, _ in sections:
                for t, lvl in pdf_parser.outlines:
                    tks = set([t[i] + t[i + 1] for i in range(len(t) - 1)])
                    tks_ = set([txt[i] + txt[i + 1] for i in range(min(len(t), len(txt) - 1))])
                    if len(set(tks & tks_)) / max([len(tks), len(tks_), 1]) > 0.8:
                        levels.append(lvl)
                        break
                else:
                    levels.append(max_lvl + 1)

        else:
            bull = bullets_category([txt for txt, _, _ in sections])
            most_level, levels = title_frequency(bull, [(txt, lvl) for txt, lvl, _ in sections])

        assert len(sections) == len(levels)
        sec_ids = []
        sid = 0
        for i, lvl in enumerate(levels):
            if lvl <= most_level and i > 0 and lvl != levels[i - 1]:
                sid += 1
            sec_ids.append(sid)

        sections = [(txt, sec_ids[i], poss) for i, (txt, _, poss) in enumerate(sections)]
        for (img, rows), poss in tbls:
            if not rows:
                continue
            sections.append((rows if isinstance(rows, str) else rows[0], -1, [(p[0] + 1 - from_page, p[1], p[2], p[3], p[4]) for p in poss]))

        def tag(pn, left, right, top, bottom):
            if pn + left + right + top + bottom == 0:
                return ""
            return "@@{}\t{:.1f}\t{:.1f}\t{:.1f}\t{:.1f}##".format(pn, left, right, top, bottom)

        chunks = []
        last_sid = -2
        tk_cnt = 0

        # Check if we're using VisionParser (all sections have level=0)
        using_vision_parser = all(sec_id == 0 for _, sec_id, _ in sections) if sections else False

        for txt, sec_id, poss in sorted(sections, key=lambda x: (x[-1][0][0], x[-1][0][3], x[-1][0][1])):
            poss = "\t".join([tag(*pos) for pos in poss])

            # For VisionParser, keep each page as separate chunk
            # For other parsers, use the original merging logic
            should_merge = False
            if not using_vision_parser:
                should_merge = tk_cnt < 32 or (tk_cnt < 1024 and (sec_id == last_sid or sec_id == -1))

            if should_merge and chunks:
                chunks[-1] += "\n" + txt + poss
                tk_cnt += num_tokens_from_string(txt)
            else:
                chunks.append(txt + poss)
                tk_cnt = num_tokens_from_string(txt)
                if sec_id > -1:
                    last_sid = sec_id

        res = tokenize_table(tbls, doc, eng)
        res.extend(tokenize_chunks(chunks, doc, eng, pdf_parser))
        return res

    elif re.search(r"\.docx?$", filename, re.IGNORECASE):
        # Try to use vision model for DOCX image processing if available
        vision_model = None
        vision_prompt = None
        layout_recognizer = parser_config.get("layout_recognize", "DeepDOC")

        if layout_recognizer not in ["DeepDOC", "Plain Text"]:
            try:
                vision_model = LLMBundle(kwargs["tenant_id"], LLMType.IMAGE2TEXT, llm_name=layout_recognizer, lang=lang)
                vision_prompt = """Please extract the layout information from the image and output only the text content in Markdown format.

1. Ignore bounding boxes and coordinates.
2. Categories:
   - Picture: omit the text field.
   - Formula: output as LaTeX format (enclosed in $ or $$).
   - Table: MUST output as HTML table format with <table>, <tr>, <td>, <th> tags. Preserve all data and structure.
   - All other categories (Text, Title, Caption, etc.): output as Markdown.

3. Constraints:
   - The output text must be the original text from the image, with no translation.
   - All layout elements must be sorted according to human reading order.
   - Tables should be formatted as complete HTML tables without unnecessary spaces or newlines.

4. Final Output: A single Markdown document containing the extracted content with tables in HTML format."""
                callback(0.05, f"Using {layout_recognizer} for image processing in DOCX.")
            except Exception as e:
                logging.info(f"Vision model {layout_recognizer} not available for DOCX: {e}")
                vision_model = None
                vision_prompt = None

        docx_parser = Docx()
        ti_list, tbls = docx_parser(filename, binary, from_page=0, to_page=10000, callback=callback, vision_model=vision_model, vision_prompt=vision_prompt)
        res = tokenize_table(tbls, doc, eng)
        for text, image in ti_list:
            d = copy.deepcopy(doc)
            if image:
                d["image"] = image
                d["doc_type_kwd"] = "image"
            tokenize(d, text, eng)
            res.append(d)
        return res
    else:
        raise NotImplementedError("file type not supported yet(pdf and docx supported)")


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], callback=dummy)
